import streamlit as st  #Webアプリのライブラリ
import fitz #PyMuPDF:PDF操作用
from google import genai
from docx import Document   #Word作成用
import time #待ち時間用
import os
#from dotenv import load_dotenv
import random
import concurrent.futures
import requests
import re
import textwrap

def get_gemini_models(api_key):
    try:
        # モデル一覧を取得するエンドポイント
        list_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
        res = requests.get(list_url)
        if res.status_code == 200:
            data = res.json()
            # generateContent（テキスト生成）に対応しているモデルのみを抽出
            models = [
                m["name"] for m in data.get("models", []) 
                if "generateContent" in m.get("supportedGenerationMethods", [])
            ]
            return models
        else:
            return []
    except Exception:
        return []

#load_dotenv()
#client = genai.Client(api_key=os.getenv("GEMINI_API_KEY")) #.envファイルからAPIキーを取得
api_key=st.secrets["GEMINI_API_KEY"]
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"]) #st.secretsからAPIキーを取得

# セッション状態（データの記憶）の初期化
if "result_path" not in st.session_state:
    st.session_state.result_path = None
if "final_filename" not in st.session_state:
    st.session_state.final_filename = None
if "is_processing" not in st.session_state:
    st.session_state.is_processing = False

st.title("PDF翻訳") #アプリのタイトル

def translate_worker(index, chunk, model_name):
    max_retries = 10
    for attempt in range(max_retries):
        try:
            prompt = textwrap.dedent(f"""
                                    # Role
                                    あなたは学術論文の翻訳を専門とする、プロフェッショナルな翻訳家です。
                                    情報通信工学、人間工学、および行動心理学の高度な専門知識を持ち、文脈に応じた適切な専門用語の選定と、流暢で自然な日本語表現を両立させるエキスパートです。
                                    
                                    # Task
                                    入力される英語の論文内容を、以下の【厳守ルール】に従って日本語に翻訳してください。
                                    # 【厳守ルール】
                                    1. **完全翻訳の徹底**: 
                                    - 入力されたテキストに含まれる全ての文章を、一文も漏らさず、省略することなく全て翻訳してください。
                                    - 要約や「以下省略」といった処理は絶対に禁止します。
                                    2. **専門用語の最適化**: 
                                    - 文脈（特に情報通信や人間工学）に合わせ、日本の学会や論文で一般的に用いられる標準的な専門用語を適切に選択してください。
                                    3. **出力形式の制限**: 
                                    - **翻訳した内容のみ**を出力してください。
                                    - 「はい、翻訳しました」「以下が翻訳結果です」といった挨拶、前置き、解説などは一切含めないでください。
                                    4. **エラーハンドリング**: 
                                    - 入力内容が極端に不完全である、または技術的に翻訳が極めて困難であると判断した場合は、理由を述べず **「**翻訳エラー**」** とだけ返してください。

                                    # 翻訳対象のテキスト
                                    {chunk}
                                    """).strip()

            response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
            return index, response.text #成功で順番と翻訳結果を返す
        except Exception as e:
            if "503" in str(e) or "429" in str(e):
                wait_time = (2 ** attempt) + random.uniform(0,1) + 5
                time.sleep(wait_time)
            else:
                return index, f"【エラー発生】：{str(e)}"   #致命的なエラーはあきらめる
    return index,   "【リトライ回数を超過しました】"

def translate_pdf_parallel(pdf_file, model_name, doc_name):
    pdf_file.seek(0)
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf") #pdfを開く
    full_text = ""
    for page in doc:
        #blocks=Trueにすることで、論文での二段組みなどの構造をある程度考慮して読み取る
        blocks = page.get_text("blocks")
        for b in blocks:
            full_text += b[4] + "\n"    #テキスト内容をfull_textに保持

    doc.close()

    chunk_size = 6000
    chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
    results = [None] * len(chunks)  #結果を格納する空のリスト

    progress_text = "並列翻訳実行中..."
    my_bar = st.progress(0, text=progress_text)

    ## ここでmax_workerを定義して何回並列に回すかを決定する
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        #処理を予約
        future_to_idx = {executor.submit(translate_worker, i, chunk, model_name): i for i, chunk in enumerate(chunks)}

        #完了したものから順に受け取る
        completed_count = 0
        for future in concurrent.futures.as_completed(future_to_idx):
            idx, translated_part = future.result() #結果を取得
            results[idx] = translated_part  #元の順番に格納

            completed_count += 1
            progress = completed_count / len(chunks)    #進捗率計算
            my_bar.progress(progress,text=f"{progress_text} ({completed_count}/{len(chunks)} 完了)")
    
    #すべての翻訳結果を結合
    full_translated_text = "\n\n".join(results)

    #Wordファイルとして保存
    output_doc = Document()
    output_doc.add_heading(doc_name,0) #見出しを追加
    for line in full_translated_text.split("\n"):
        if line.strip():
            output_doc.add_paragraph(line)

    safe_name = re.sub(r'[\\/:\*\?"<>\|]', '_', doc_name)
    if not safe_name.strip():
        safe_name = "translated_result"

    output_filename = f"{safe_name}.docx"
    output_doc.save(output_filename)    #ファイルを保存
    return output_filename

uploaded_file = st.file_uploader("PDFファイルを選択していください", type="pdf") #ファイルの選択

doc_name = st.text_input("Wordタイトル", value="論文和訳結果")

target_model = None
if api_key:
    with st.spinner("利用可能なGeminiモデルを確認中..."):
        available_models = get_gemini_models(api_key)
    
    if available_models:
        # デフォルトで「gemini-3.1-flash-lite」や「gemini-2.5-flash」などがあればそれを初期値にするロジック
        default_index = 0
        for i, m in enumerate(available_models):
            if "gemini-3.1-flash-lite" in m:
                default_index = i
                break
        
        target_model = st.selectbox(
            "使用するAIモデルを選択してください（APIから自動取得）",
            options=available_models,
            index=default_index
        )
    else:
        st.warning("モデル一覧の取得に失敗しました。デフォルトのモデルを使用します。")
        target_model = "models/gemini-3.1-flash-lite"
else:
    st.error("APIキーが設定されていないため、モデル一覧を取得できません。")
if uploaded_file is not None:
    st.info("ファイルがアップロードされました。下のボタンを押すと翻訳を開始します。")   
    if st.button("翻訳を開始",disabled=st.session_state.is_processing):
        st.session_state.is_processing = True

        with st.spinner("AIが思考中..."):
            try:
                st.session_state.result_path = translate_pdf_parallel(uploaded_file, target_model, doc_name)
                st.session_state.final_filename = os.path.basename(st.session_state.result_path)
                st.success("翻訳が完了しました！")
            except Exception as e:
                st.error(f"エラーが発生しました: {e}")
            finally:
                st.session_state.is_processing = False
                #st.rerun()

#翻訳が完了している場合のみダウンロードボタンを表示し続ける
if st.session_state.result_path and os.path.exists(st.session_state.result_path):
    st.divider()   
# --- 5. ダウンロードボタンを設置 ---
    with open(st.session_state.result_path, "rb") as f:
        st.download_button(
            label="翻訳後のWordファイルをダウンロード",
            data=f,
            file_name=st.session_state.final_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
            